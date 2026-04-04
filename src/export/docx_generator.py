from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches
from src.i18n.languages import is_rtl
from src.logger import get_logger

logger = get_logger(__name__)


class DocxGenerator:
    def __init__(self, projects_dir: Path | str = "projects"):
        self.projects_dir = Path(projects_dir)

    def generate(
        self,
        project_id: int,
        title: str = "Ebook",
        subtitle: str = "",
        author: str = "Author",
        language: str = "en",
    ) -> dict:
        project_dir = self.projects_dir / str(project_id)
        exports_dir = project_dir / "exports"
        exports_dir.mkdir(parents=True, exist_ok=True)

        doc = Document()

        self._add_cover_page(doc, title, subtitle, author, project_dir)

        doc.add_page_break()

        self._add_copyright_page(doc)

        self._add_toc(doc, project_id)

        doc.add_page_break()

        manuscript_file = project_dir / "manuscript.md"
        if manuscript_file.exists():
            self._parse_manuscript_to_docx(doc, manuscript_file, language)

        self._apply_styles(doc)

        docx_file = exports_dir / "ebook.docx"
        doc.save(docx_file)

        return {"docx": docx_file}

    def _add_cover_page(
        self, doc: Document, title: str, subtitle: str, author: str, project_dir: Path
    ) -> None:
        cover_file = project_dir / "cover" / "cover.png"

        if cover_file.exists():
            try:
                doc.add_picture(str(cover_file), width=Inches(6))
            except Exception as e:
                logger.warning("Failed to add cover picture to DOCX", error=str(e))

        doc.add_heading(title, 0)
        if subtitle:
            doc.add_paragraph(subtitle)
        doc.add_paragraph(f"By {author}")

    def _add_copyright_page(self, doc: Document) -> None:
        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph(f"Copyright © {datetime.now().year}")
        doc.add_paragraph("All rights reserved.")

    def _add_toc(self, doc: Document, project_id: int) -> None:
        doc.add_heading("Table of Contents", 1)

        toc_file = self.projects_dir / str(project_id) / "toc.md"

        if toc_file.exists():
            content = toc_file.read_text()
            for line in content.split("\n"):
                if line.strip().startswith("## ") and "Table of Contents" not in line:
                    title = line.replace("## ", "").strip()
                    doc.add_paragraph(title)

    def _set_rtl(self, paragraph) -> None:
        """Apply RTL text direction to a paragraph using OOXML w:bidi element."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        pPr.append(bidi)

    def _parse_manuscript_to_docx(self, doc: Document, manuscript_file: Path, language: str = "en") -> None:
        content = manuscript_file.read_text()

        lines = content.split("\n")
        chapter_index = 0

        for line in lines:
            line = line.strip()

            if line.startswith("# "):
                title = line.replace("# ", "").strip()
                heading = doc.add_heading(title, 1)
                heading.paragraph_format.page_break_before = chapter_index > 0
                chapter_index += 1
                if is_rtl(language):
                    self._set_rtl(heading)

            elif line.startswith("## "):
                title = line.replace("## ", "").strip()
                doc.add_heading(title, 2)
                heading_para = doc.paragraphs[-1]
                if is_rtl(language):
                    self._set_rtl(heading_para)

            elif line and not line.startswith("#"):
                para = doc.add_paragraph(line)
                if is_rtl(language):
                    self._set_rtl(para)

    def _apply_styles(self, doc: Document) -> None:
        from docx.shared import Pt

        # Body text
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)

        # Heading 1
        h1 = doc.styles["Heading 1"]
        h1.font.name = "Calibri"
        h1.font.size = Pt(18)
        h1.font.bold = True

        # Heading 2
        h2 = doc.styles["Heading 2"]
        h2.font.name = "Calibri"
        h2.font.size = Pt(14)
        h2.font.bold = True

        # Heading 3
        try:
            h3 = doc.styles["Heading 3"]
            h3.font.name = "Calibri"
            h3.font.size = Pt(12)
            h3.font.bold = True
            h3.font.italic = True
        except KeyError:
            pass
