import pytest
import subprocess
from unittest.mock import patch, MagicMock
from pathlib import Path


def test_convert_docx_to_pdf(temp_project_dir):
    from src.export.pdf_converter import PdfConverter

    project_dir = temp_project_dir / "1"
    project_dir.mkdir(parents=True, exist_ok=True)
    (project_dir / "exports").mkdir(exist_ok=True)

    docx_file = project_dir / "exports" / "ebook.docx"
    from docx import Document

    doc = Document()
    doc.add_heading("Test", 0)
    doc.add_paragraph("Content")
    doc.save(docx_file)

    pdf_file = project_dir / "exports" / "ebook.pdf"
    pdf_file.write_bytes(b"%PDF-1.4 fake pdf content")

    converter = PdfConverter(libreoffice_path="/usr/bin/libreoffice", projects_dir=temp_project_dir)
    with patch("subprocess.run") as mock_run:
        mock_run.return_value = MagicMock(returncode=0, stderr="")
        result = converter.convert(docx_file)

    assert pdf_file.exists()


def test_pdf_starts_with_pdf_header(temp_project_dir):
    from src.export.pdf_converter import PdfConverter

    project_dir = temp_project_dir / "1"
    project_dir.mkdir(parents=True, exist_ok=True)
    (project_dir / "exports").mkdir(exist_ok=True)

    docx_file = project_dir / "exports" / "ebook.docx"
    from docx import Document

    doc = Document()
    doc.add_paragraph("Test")
    doc.save(docx_file)

    pdf_file = project_dir / "exports" / "ebook.pdf"
    pdf_file.write_bytes(b"%PDF-1.4")

    converter = PdfConverter(libreoffice_path="/usr/bin/libreoffice", projects_dir=temp_project_dir)
    with patch("subprocess.run") as mock_run:
        mock_run.return_value = MagicMock(returncode=0, stderr="")
        converter.convert(docx_file)

    content = pdf_file.read_bytes()[:4]
    assert content == b"%PDF"


def test_error_when_libreoffice_missing(temp_project_dir):
    from src.export.pdf_converter import PdfConverter

    project_dir = temp_project_dir / "1"
    project_dir.mkdir(parents=True, exist_ok=True)
    (project_dir / "exports").mkdir(exist_ok=True)

    docx_file = project_dir / "exports" / "ebook.docx"
    docx_file.touch()

    converter = PdfConverter(projects_dir=temp_project_dir)

    with patch("subprocess.run", side_effect=FileNotFoundError(" libreoffice")):
        with pytest.raises(RuntimeError, match="LibreOffice not found"):
            converter.convert(docx_file)


def test_check_installation_returns_true_when_libreoffice_found(temp_project_dir):
    from src.export.pdf_converter import PdfConverter

    converter = PdfConverter(libreoffice_path="/usr/bin/libreoffice", projects_dir=temp_project_dir)
    assert converter.check_installation() is True


def test_check_installation_returns_false_when_libreoffice_not_found(temp_project_dir):
    from src.export.pdf_converter import PdfConverter

    converter = PdfConverter(libreoffice_path=None, projects_dir=temp_project_dir)
    with patch("shutil.which", return_value=None):
        converter = PdfConverter(projects_dir=temp_project_dir)
        assert converter.check_installation() is False


def test_conversion_timeout_raises_error(temp_project_dir):
    from src.export.pdf_converter import PdfConverter

    project_dir = temp_project_dir / "1"
    project_dir.mkdir(parents=True, exist_ok=True)
    (project_dir / "exports").mkdir(exist_ok=True)

    docx_file = project_dir / "exports" / "ebook.docx"
    from docx import Document
    doc = Document()
    doc.add_paragraph("Test")
    doc.save(docx_file)

    converter = PdfConverter(libreoffice_path="/usr/bin/libreoffice", projects_dir=temp_project_dir)

    with patch("subprocess.run", side_effect=subprocess.TimeoutExpired("libreoffice", 120)):
        with pytest.raises(subprocess.TimeoutExpired):
            converter.convert(docx_file)


def test_subprocess_failure_raises_runtime_error(temp_project_dir):
    from src.export.pdf_converter import PdfConverter

    project_dir = temp_project_dir / "1"
    project_dir.mkdir(parents=True, exist_ok=True)
    (project_dir / "exports").mkdir(exist_ok=True)

    docx_file = project_dir / "exports" / "ebook.docx"
    from docx import Document
    doc = Document()
    doc.add_paragraph("Test")
    doc.save(docx_file)

    converter = PdfConverter(libreoffice_path="/usr/bin/libreoffice", projects_dir=temp_project_dir)

    with patch("subprocess.run") as mock_run:
        mock_run.return_value = MagicMock(returncode=1, stderr="Conversion error: corrupted file")
        with pytest.raises(RuntimeError, match="PDF conversion failed: Conversion error"):
            converter.convert(docx_file)


def test_path_traversal_attack_blocked(temp_project_dir):
    from src.export.pdf_converter import PdfConverter

    project_dir = temp_project_dir / "1"
    project_dir.mkdir(parents=True, exist_ok=True)
    (project_dir / "exports").mkdir(exist_ok=True)

    outside_dir = temp_project_dir.parent / "outside"
    outside_dir.mkdir(exist_ok=True)
    malicious_file = outside_dir / "malicious.docx"
    malicious_file.touch()

    converter = PdfConverter(libreoffice_path="/usr/bin/libreoffice", projects_dir=temp_project_dir)

    with pytest.raises(ValueError, match="Invalid file path"):
        converter.convert(malicious_file)


def test_symlink_traversal_blocked(temp_project_dir):
    from src.export.pdf_converter import PdfConverter

    project_dir = temp_project_dir / "1"
    project_dir.mkdir(parents=True, exist_ok=True)
    (project_dir / "exports").mkdir(exist_ok=True)

    outside_dir = temp_project_dir.parent / "outside"
    outside_dir.mkdir(exist_ok=True)
    target_file = outside_dir / "secret.docx"
    target_file.touch()

    symlink_file = project_dir / "exports" / "link.docx"
    symlink_file.symlink_to(target_file)

    converter = PdfConverter(libreoffice_path="/usr/bin/libreoffice", projects_dir=temp_project_dir)

    with pytest.raises(ValueError, match="Invalid file path"):
        converter.convert(symlink_file)


def test_invalid_file_extension_rejected(temp_project_dir):
    from src.export.pdf_converter import PdfConverter

    project_dir = temp_project_dir / "1"
    project_dir.mkdir(parents=True, exist_ok=True)
    (project_dir / "exports").mkdir(exist_ok=True)

    txt_file = project_dir / "exports" / "ebook.txt"
    txt_file.write_text("Not a DOCX file")

    converter = PdfConverter(libreoffice_path="/usr/bin/libreoffice", projects_dir=temp_project_dir)

    with pytest.raises(ValueError, match="File must have one of these extensions"):
        converter.convert(txt_file)


def test_nonexistent_file_rejected(temp_project_dir):
    from src.export.pdf_converter import PdfConverter

    project_dir = temp_project_dir / "1"
    project_dir.mkdir(parents=True, exist_ok=True)
    (project_dir / "exports").mkdir(exist_ok=True)

    nonexistent_file = project_dir / "exports" / "nonexistent.docx"

    converter = PdfConverter(libreoffice_path="/usr/bin/libreoffice", projects_dir=temp_project_dir)

    with pytest.raises(ValueError, match="File does not exist"):
        converter.convert(nonexistent_file)


def test_special_characters_in_path_handled(temp_project_dir):
    from src.export.pdf_converter import PdfConverter

    project_dir = temp_project_dir / "1"
    project_dir.mkdir(parents=True, exist_ok=True)
    (project_dir / "exports").mkdir(exist_ok=True)

    docx_file = project_dir / "exports" / "ebook with spaces & special.docx"
    from docx import Document
    doc = Document()
    doc.add_paragraph("Test")
    doc.save(docx_file)

    pdf_file = project_dir / "exports" / "ebook with spaces & special.pdf"
    pdf_file.write_bytes(b"%PDF-1.4 fake pdf content")

    converter = PdfConverter(libreoffice_path="/usr/bin/libreoffice", projects_dir=temp_project_dir)

    with patch("subprocess.run") as mock_run:
        mock_run.return_value = MagicMock(returncode=0, stderr="")
        result = converter.convert(docx_file)

    assert result["pdf"] == pdf_file
    assert pdf_file.exists()


def test_pdf_not_created_raises_error(temp_project_dir):
    from src.export.pdf_converter import PdfConverter

    project_dir = temp_project_dir / "1"
    project_dir.mkdir(parents=True, exist_ok=True)
    (project_dir / "exports").mkdir(exist_ok=True)

    docx_file = project_dir / "exports" / "ebook.docx"
    from docx import Document
    doc = Document()
    doc.add_paragraph("Test")
    doc.save(docx_file)

    converter = PdfConverter(libreoffice_path="/usr/bin/libreoffice", projects_dir=temp_project_dir)

    with patch("subprocess.run") as mock_run:
        mock_run.return_value = MagicMock(returncode=0, stderr="")
        with pytest.raises(RuntimeError, match="PDF file not created"):
            converter.convert(docx_file)


def test_case_insensitive_extension_validation(temp_project_dir):
    from src.export.pdf_converter import PdfConverter

    project_dir = temp_project_dir / "1"
    project_dir.mkdir(parents=True, exist_ok=True)
    (project_dir / "exports").mkdir(exist_ok=True)

    docx_file = project_dir / "exports" / "ebook.DOCX"
    from docx import Document
    doc = Document()
    doc.add_paragraph("Test")
    doc.save(docx_file)

    pdf_file = project_dir / "exports" / "ebook.pdf"
    pdf_file.write_bytes(b"%PDF-1.4 fake pdf content")

    converter = PdfConverter(libreoffice_path="/usr/bin/libreoffice", projects_dir=temp_project_dir)

    with patch("subprocess.run") as mock_run:
        mock_run.return_value = MagicMock(returncode=0, stderr="")
        result = converter.convert(docx_file)

    assert result["pdf"] == pdf_file


def test_subprocess_command_uses_validated_path(temp_project_dir):
    from src.export.pdf_converter import PdfConverter

    project_dir = temp_project_dir / "1"
    project_dir.mkdir(parents=True, exist_ok=True)
    (project_dir / "exports").mkdir(exist_ok=True)

    docx_file = project_dir / "exports" / "ebook.docx"
    from docx import Document
    doc = Document()
    doc.add_paragraph("Test")
    doc.save(docx_file)

    pdf_file = project_dir / "exports" / "ebook.pdf"
    pdf_file.write_bytes(b"%PDF-1.4 fake pdf content")

    converter = PdfConverter(libreoffice_path="/usr/bin/libreoffice", projects_dir=temp_project_dir)

    with patch("subprocess.run") as mock_run:
        mock_run.return_value = MagicMock(returncode=0, stderr="")
        converter.convert(docx_file)

    mock_run.assert_called_once()
    call_args = mock_run.call_args[0][0]
    assert call_args[0] == "/usr/bin/libreoffice"
    assert call_args[1] == "--headless"
    assert call_args[2] == "--convert-to"
    assert call_args[3] == "pdf"
    assert call_args[4] == "--outdir"
    assert str(docx_file.resolve()) == call_args[6]


def test_libreoffice_auto_detection(temp_project_dir):
    from src.export.pdf_converter import PdfConverter

    with patch("shutil.which") as mock_which:
        mock_which.side_effect = lambda cmd: "/usr/bin/libreoffice" if cmd == "libreoffice" else None
        converter = PdfConverter(projects_dir=temp_project_dir)
        assert converter.libreoffice_path == "/usr/bin/libreoffice"


def test_soffice_fallback_detection(temp_project_dir):
    from src.export.pdf_converter import PdfConverter

    with patch("shutil.which") as mock_which:
        mock_which.side_effect = lambda cmd: "/usr/bin/soffice" if cmd == "soffice" else None
        converter = PdfConverter(projects_dir=temp_project_dir)
        assert converter.libreoffice_path == "/usr/bin/soffice"
