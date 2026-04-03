import pytest
from pathlib import Path


def test_docx_generated_from_manuscript(temp_project_dir):
    from src.export.docx_generator import DocxGenerator

    project_dir = temp_project_dir / "1"
    project_dir.mkdir(parents=True, exist_ok=True)

    (project_dir / "chapters").mkdir(exist_ok=True)
    (project_dir / "chapters" / "1.md").write_text("# Chapter 1\n\nContent here.")
    (project_dir / "chapters" / "2.md").write_text("# Chapter 2\n\nMore content.")
    (project_dir / "manuscript.md").write_text(
        "# Chapter 1\n\nContent here.\n\n# Chapter 2\n\nMore content."
    )

    cover_dir = project_dir / "cover"
    cover_dir.mkdir(exist_ok=True)

    from PIL import Image

    img = Image.new("RGB", (100, 100), (255, 255, 255))
    img.save(cover_dir / "cover.png")

    generator = DocxGenerator(projects_dir=temp_project_dir)
    result = generator.generate(project_id=1, title="Test Ebook", author="Test Author")

    docx_file = project_dir / "exports" / "ebook.docx"
    assert docx_file.exists()


def test_docx_valid_zip_format(temp_project_dir):
    from src.export.docx_generator import DocxGenerator

    project_dir = temp_project_dir / "1"
    project_dir.mkdir(parents=True, exist_ok=True)
    (project_dir / "chapters").mkdir(exist_ok=True)
    (project_dir / "chapters" / "1.md").write_text("# Ch1\n\nContent.")
    (project_dir / "manuscript.md").write_text("# Ch1\n\nContent.")
    (project_dir / "cover").mkdir(exist_ok=True)

    from PIL import Image

    Image.new("RGB", (100, 100)).save(project_dir / "cover" / "cover.png")

    generator = DocxGenerator(projects_dir=temp_project_dir)
    generator.generate(project_id=1)

    docx_file = project_dir / "exports" / "ebook.docx"
    content = docx_file.read_bytes()
    assert content[:4] == b"PK\x03\x04"


def test_docx_has_h1_styles(temp_project_dir):
    from src.export.docx_generator import DocxGenerator

    project_dir = temp_project_dir / "1"
    project_dir.mkdir(parents=True, exist_ok=True)
    (project_dir / "chapters").mkdir(exist_ok=True)
    (project_dir / "chapters" / "1.md").write_text("# Chapter One\n\nContent")
    (project_dir / "manuscript.md").write_text("# Chapter One\n\nContent")
    (project_dir / "cover").mkdir(exist_ok=True)

    from PIL import Image

    Image.new("RGB", (100, 100)).save(project_dir / "cover" / "cover.png")

    generator = DocxGenerator(projects_dir=temp_project_dir)
    generator.generate(project_id=1)

    docx_file = project_dir / "exports" / "ebook.docx"
    import zipfile

    with zipfile.ZipFile(docx_file) as z:
        xml = z.read("word/document.xml").decode()
        assert "Chapter One" in xml


def test_apply_styles_sets_calibri(tmp_path):
    from docx import Document
    from src.export.docx_generator import DocxGenerator
    gen = DocxGenerator(projects_dir=tmp_path)
    doc = Document()
    gen._apply_styles(doc)
    assert doc.styles["Normal"].font.name == "Calibri"
    assert doc.styles["Heading 1"].font.name == "Calibri"


def test_set_rtl_adds_bidi_element(tmp_path):
    from docx import Document
    from src.export.docx_generator import DocxGenerator
    gen = DocxGenerator(projects_dir=tmp_path)
    doc = Document()
    para = doc.add_paragraph("مرحبا")
    gen._set_rtl(para)
    xml_str = para._p.xml
    assert "w:bidi" in xml_str
